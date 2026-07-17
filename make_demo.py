"""Build a realistic (fake) MSP escalation ticket as a .docx demo input."""

import docx
from docx.shared import Pt

doc = docx.Document()

doc.add_heading("Escalation Ticket #48213 — VPN + File Share Outage", 0)

p = doc.add_paragraph()
p.add_run("Client: ").bold = True
p.add_run("Harbor Point Dental Group\n")
p.add_run("Contact: ").bold = True
p.add_run("Sarah Whitfield — sarah.whitfield@harborpointdental.com, "
          "cell (415) 555-0173\n")
p.add_run("Technician: ").bold = True
p.add_run("Marcus Reyes (Tier 2)\n")
p.add_run("Priority: ").bold = True
p.add_run("P2 — business hours impact")

doc.add_heading("Environment", level=1)
table = doc.add_table(rows=5, cols=2)
table.style = "Light Grid Accent 1"
rows = [
    ("Domain controller", "HPD-DC01 — 192.168.20.10 (corp.harborpoint.local)"),
    ("File server", r"HPD-FS02 — 192.168.20.14, shares at \\HPD-FS02\PatientDocs"),
    ("Firewall WAN", "Static IP 52.96.114.88, mgmt MAC 00:1A:2B:3C:4D:5E"),
    ("M365 tenant", "Tenant ID 3f2b8c1d-4e5a-4f89-abcd-ef0123456789"),
    ("Backup NAS", "HPD-NAS01 — 192.168.20.30 (Synology)"),
]
for i, (k, v) in enumerate(rows):
    table.rows[i].cells[0].text = k
    table.rows[i].cells[1].text = v

doc.add_heading("Troubleshooting Notes", level=1)
doc.add_paragraph(
    "User reports VPN drops every ~20 min. Verified from HPD-DC01 that "
    "the RADIUS service was restarting. Temporarily connected to the "
    "firewall API for diagnostics using api_key = fgt_7Hq2mX9pLw4Zt8Rb3Kd6 "
    "(needs rotation after this ticket)."
)
doc.add_paragraph(
    "Storage account for offsite backup was re-linked. Connection string "
    "used during testing: DefaultEndpointsProtocol=https;AccountName=hpdbackup;"
    "AccountKey=q8Zx2Wv4Yt6Us8Qp0On2Ml4Kj6Ih8Gf0Ed2Cb4Az6xYw==;"
    "EndpointSuffix=core.windows.net"
)
doc.add_paragraph(
    "SQL maintenance job on HPD-SQL01 failed overnight — reran with "
    "Server=HPD-SQL01;Database=DentrixDB;User Id=svc_backup;"
    "Password=Sp1ng!Fl0wer#88; — confirmed success."
)
doc.add_paragraph(
    "Patient called in during the incident; verified identity with last-4 "
    "and DOB, full SSN on file reads 078-05-1120 (flagged: should never "
    "have been pasted into the ticket — compliance issue)."
)
doc.add_paragraph(
    "Approved by: Daniel Okafor. Follow-up scheduled with Sarah for Friday. "
    "Session JWT captured in HAR file: "
    "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9."
    "eyJzdWIiOiIxMjM0NTY3ODkwIiwibmFtZSI6IlNhcmFoIn0."
    "SflKxwRJSMeKKF2QT4fwpMeJf36POk6yJVadQssw5c"
)

doc.add_heading("Next Steps", level=1)
doc.add_paragraph(
    "1. Rotate firewall API key and SQL svc_backup password.\n"
    "2. Review VLAN 20 DHCP scope on 192.168.20.0/24.\n"
    "3. Purge SSN from ticket history per HIPAA policy.\n"
    "4. Email RCA to sarah.whitfield@harborpointdental.com and cc "
    "m.reyes@ourmsp.com."
)

doc.save("demo/escalation_ticket_48213.docx")
print("demo doc written")
