"""Test doc #2 — harder than the demo: names without context words (NER
stress), more secret formats, tricky negatives that should NOT be redacted."""

import docx

doc = docx.Document()
doc.add_heading("Network Migration Runbook — Cutover Weekend", 0)

doc.add_paragraph(
    "Scope: migrate Bayview Logistics from the legacy Meraki stack to "
    "UniFi. Cutover window approved during Thursday's call. Elena Vasquez "
    "will be onsite; Tom Abernathy covers remote hands. Escalation goes "
    "through Priya Raman if the change runs past 02:00."
)

doc.add_heading("Current Environment", level=1)
doc.add_paragraph(
    "Core switch BVL-SW01 (10.44.0.2) uplinks to the edge at 10.44.0.1. "
    "Management VLAN is 10.44.99.0/24 with the controller on "
    "vc.bayviewlog.corp (10.44.99.5). Wireless bridge MAC is "
    "E4:38:83:1B:7C:2F. Legacy AP mesh uses IPv6 ULA fd00:44:0:99::1. "
    "Public egress NAT: 40.117.83.202. Subnet mask 255.255.255.0, "
    "gateway metric 10, MTU 1500, port 8443."
)

doc.add_heading("Credentials Staged for Cutover", level=1)
doc.add_paragraph(
    "UniFi controller admin created; recovery seed stored in Hudu. "
    "Cloud key API token: AIzaSyD-9tSrke72PouQMnMX-a7eZSW0jkFMBWY (rotate "
    "post-cutover). GitHub deploy token for the config repo: "
    "ghp_wJ8kQmN4vX2rT9bL6cY3fH1sD5aG7pZeK0uM. Slack webhook auth "
    "xoxb-2847391056-KjH8gF3dS2aQ9wE4rT6yU1iO. RADIUS shared secret = "
    "Tr0ub4dor&3-horse-battery."
)
doc.add_paragraph(
    "AWS route53 automation user AKIAIOSFODNN7EXAMPLE is scoped to the "
    "bayviewlog.com zone only. S3 sync job connection: "
    "Server=BVL-SQL02;Database=WMS;User Id=wms_ro;Pwd=Gr@nite!Fj0rd42;"
)
doc.add_paragraph("Firewall management cert (staging only):")
doc.add_paragraph(
    "-----BEGIN RSA PRIVATE KEY-----\n"
    "MIIEpAIBAAKCAQEA7c4T9Xq2mZ8vN1bK5jW3rY6uP0oL4iH8gF2dS9aQ1wE5rT7y\n"
    "U3iO6pA8sD0fG2hJ4kL6zX8cV0bN2mQ4wE6rT8yU0iO2pA4sD6fG8hJ0kL2zX4cV\n"
    "-----END RSA PRIVATE KEY-----"
)

doc.add_heading("Contacts & Billing", level=1)
doc.add_paragraph(
    "Invoice questions go to accounts@bayviewlogistics.com or call "
    "(628) 555-0142. The owner keeps a card on file ending in the number "
    "4532 7597 3454 8801 for emergency hardware buys. Warehouse manager "
    "Dwayne Carter Jr. signs off on downtime windows; his cell is "
    "628.555.0177."
)

doc.add_heading("Reference (do not redact)", level=1)
doc.add_paragraph(
    "Firmware 7.0.25, Windows Server 2022, KB5034441, RFC 1918, "
    "UniFi OS 3.2.12, ticket #77201, PO 2026-0448, driver 31.0.15.5222. "
    "Change approved 2026-07-15. Version 10.0.19045.4046 build noted "
    "during the Defender audit."
)

doc.save("demo/migration_runbook_bvl.docx")
print("written")
