"""File format handlers: txt, docx, pdf.

docx: paragraphs, tables, headers/footers are sanitized in place on a copy.
PoC limitation — a modified paragraph keeps paragraph-level styling but
collapses per-run (character-level) formatting inside that one paragraph.
Untouched paragraphs are byte-identical.

pdf: text is extracted and sanitized; output is a .txt (safe-share copy).
Rebuilding a visually identical sanitized PDF is a production feature.
"""

import shutil
from pathlib import Path


# ---------------------------------------------------------------- txt

def sanitize_txt(scrubber, in_path, out_path):
    text = Path(in_path).read_text(encoding="utf-8", errors="replace")
    sanitized, findings = scrubber.sanitize_text(text)
    Path(out_path).write_text(sanitized, encoding="utf-8")
    return findings


# ---------------------------------------------------------------- docx

def _iter_paragraphs(doc):
    from docx.table import Table
    for para in doc.paragraphs:
        yield para
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    yield para
                for tbl in cell.tables:  # nested tables
                    for r in tbl.rows:
                        for c in r.cells:
                            for p in c.paragraphs:
                                yield p
    for section in doc.sections:
        for hf in (section.header, section.footer):
            for para in hf.paragraphs:
                yield para


def _replace_paragraph_text(para, new_text):
    if para.runs:
        para.runs[0].text = new_text
        for run in para.runs[1:]:
            run.text = ""
    else:
        para.add_run(new_text)


def sanitize_docx(scrubber, in_path, out_path):
    import docx
    shutil.copyfile(in_path, out_path)
    doc = docx.Document(out_path)
    all_findings = []
    for para in _iter_paragraphs(doc):
        if not para.text.strip():
            continue
        sanitized, findings = scrubber.sanitize_text(para.text)
        if findings:
            _replace_paragraph_text(para, sanitized)
            all_findings.extend(findings)
    doc.save(out_path)
    return all_findings


# ---------------------------------------------------------------- xlsx

def sanitize_xlsx(scrubber, in_path, out_path):
    """Sanitize string cells across all worksheets, preserving styling.

    Formula cells are left untouched (rewriting them corrupts references).
    Caveats (PoC): charts/images embedded in the workbook may not survive
    the openpyxl round-trip; sheet names are not renamed (formula refs).
    """
    import openpyxl

    wb = openpyxl.load_workbook(in_path)
    all_findings = []

    # Structure-aware: a column whose header says "name" holds names —
    # register every string cell under it as a known PERSON, so detection
    # doesn't depend on an NER model judging bare two-word cells.
    name_headers = {"name", "display name", "displayname", "full name",
                    "fullname", "user", "username", "owner", "contact",
                    "technician", "manager", "requester", "approver",
                    "first name", "last name", "employee", "customer"}
    for ws in wb.worksheets:
        first = next(ws.iter_rows(min_row=1, max_row=1), None) or []
        for cell in first:
            header = str(cell.value or "").strip().lower()
            if header in name_headers:
                for (val,) in ws.iter_rows(min_row=2, min_col=cell.column,
                                           max_col=cell.column,
                                           values_only=True):
                    if isinstance(val, str) and val.strip():
                        scrubber.add_known_value("PERSON", val)

    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                v = cell.value
                if not isinstance(v, str) or not v.strip():
                    continue
                if v.startswith("="):        # formula — leave alone
                    continue
                sanitized, findings = scrubber.sanitize_text(v)
                if findings:
                    cell.value = sanitized
                    all_findings.extend(findings)
    wb.save(out_path)
    return all_findings


# ---------------------------------------------------------------- email

def sanitize_eml(scrubber, in_path, out_path):
    """Sanitize an RFC-822 .eml: address headers, subject, and text bodies.

    Structure, attachments, and non-text parts pass through unchanged
    (attachment contents are NOT scanned — noted in the report via findings
    only covering text parts)."""
    import email
    import email.policy
    import email.utils

    msg = email.message_from_bytes(Path(in_path).read_bytes(),
                                   policy=email.policy.default)
    all_findings = []

    # Structure-aware: a display name in an address header IS a person —
    # register it so it's redacted deterministically everywhere (headers,
    # greeting lines, signatures), no NER required.
    for header in ("From", "To", "Cc", "Bcc", "Reply-To"):
        values = msg.get_all(header) or []
        for name, _addr in email.utils.getaddresses([str(v) for v in values]):
            if name:
                scrubber.add_known_value("PERSON", name)

    for header in ("From", "To", "Cc", "Bcc", "Reply-To", "Subject"):
        if msg.get(header):
            sanitized, findings = scrubber.sanitize_text(str(msg[header]))
            if findings:
                msg.replace_header(header, sanitized)
                all_findings.extend(findings)

    for part in msg.walk():
        if part.get_content_maintype() != "text":
            continue
        try:
            body = part.get_content()
        except Exception:
            continue
        sanitized, findings = scrubber.sanitize_text(body)
        if findings:
            part.set_content(sanitized, subtype=part.get_content_subtype())
            all_findings.extend(findings)

    Path(out_path).write_bytes(msg.as_bytes())
    return all_findings


# Outlook .msg is an OLE compound file; text lives in well-known MAPI
# property streams. Minimal pure-python extraction (olefile) — PoC outputs
# sanitized text, not a rebuilt .msg.
_MSG_STREAMS = [
    ("Subject", "__substg1.0_0037001F"),
    ("Sender", "__substg1.0_0C1A001F"),
    ("Sender email", "__substg1.0_5D01001F"),
    ("To", "__substg1.0_0E04001F"),
    ("Cc", "__substg1.0_0E03001F"),
    ("Body", "__substg1.0_1000001F"),
]


def sanitize_msg(scrubber, in_path, out_path):
    import olefile

    ole = olefile.OleFileIO(str(in_path))
    parts = []
    try:
        for label, stream in _MSG_STREAMS:
            if ole.exists(stream):
                raw = ole.openstream(stream).read()
                value = raw.decode("utf-16-le", errors="replace").strip("\x00")
                if value.strip():
                    parts.append(f"{label}: {value}" if label != "Body"
                                 else f"\n{value}")
    finally:
        ole.close()

    text = "\n".join(parts)
    sanitized, findings = scrubber.sanitize_text(text)
    Path(out_path).write_text(sanitized, encoding="utf-8")
    return findings


# ---------------------------------------------------------------- pdf

def sanitize_pdf(scrubber, in_path, out_path):
    from pypdf import PdfReader
    reader = PdfReader(in_path)
    pages = [(page.extract_text() or "") for page in reader.pages]
    text = "\n\n".join(f"--- page {i+1} ---\n{t}" for i, t in enumerate(pages))
    sanitized, findings = scrubber.sanitize_text(text)
    Path(out_path).write_text(sanitized, encoding="utf-8")
    return findings


# ---------------------------------------------------------------- dispatch

HANDLERS = {
    ".docx": (sanitize_docx, ".docx"),
    ".xlsx": (sanitize_xlsx, ".xlsx"),
    ".xlsm": (sanitize_xlsx, ".xlsm"),
    ".eml": (sanitize_eml, ".eml"),
    ".msg": (sanitize_msg, ".txt"),   # PoC: Outlook msg sanitizes to text
    ".pdf": (sanitize_pdf, ".txt"),   # PoC: PDF sanitizes to text
}

# Plain-text family — configs, code, exports all ride the txt handler and
# come back as themselves.
for _ext in (".txt", ".md", ".log", ".csv", ".tsv", ".json", ".yaml", ".yml",
             ".xml", ".html", ".htm", ".ini", ".conf", ".cfg", ".toml",
             ".ps1", ".psm1", ".sh", ".bat", ".cmd", ".py", ".sql", ".rtf"):
    HANDLERS[_ext] = (sanitize_txt, _ext)


def sanitize_file(scrubber, in_path, out_dir):
    in_path = Path(in_path)
    ext = in_path.suffix.lower()
    if ext not in HANDLERS:
        raise ValueError(f"Unsupported file type: {ext} "
                         f"(supported: {', '.join(HANDLERS)})")
    handler, out_ext = HANDLERS[ext]
    out_dir = Path(out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / f"{in_path.stem}_SANITIZED{out_ext}"
    findings = handler(scrubber, in_path, out_path)
    return out_path, findings
